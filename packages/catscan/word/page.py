from docx.shared import Inches, Mm, Twips
from styles import check_style
# from docutils.doc import AbstractNotFoundError

class TrackingOnError(Exception):
    """Raised when the docx has tracking change on"""
    pass


AUTHOR_DETAILS = {
    'styles': {
        'jacow': 'JACoW_Author List',
        'normal': 'Author List',
    },
    'alignment': 'CENTER',
    'font_size': 12.0,
    'space_before': 9.0,
    'space_after': 12.0,
    'bold': None,
    'italic': None,
}

ABSTRACT_DETAILS = {
    'styles': {
        'jacow': 'JACoW_Abstract_Heading',
        'normal': 'Abstract_Heading',
    },
    'alignment': None,
    'font_size': 12.0,
    'space_before': 0.0,
    'space_after': 3.0,
    'bold': None,
    'italic': True,
}


def get_page_size(section):
    width = round(section.page_width, -4)
    if width == round(Mm(210), -4):
        return 'A4'
    elif width == round(Inches(8.5), -4):
        return 'Letter'
    else:
        mm = round(round(width, -4) / 36000, 2)
        return f'Unknown page size with width {mm} mm'


def get_abstract_and_author(doc):
    abstract = {}
    title_start = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() and title_start == -1:
            title_start = i

        if p.text.strip().lower() == 'abstract':
            style_ok, detail = check_style(p, ABSTRACT_DETAILS)
            abstract = {
                'start': i,
                'text': p.text,
                'style': p.style.name,
                'style_ok': style_ok,
            }
            abstract.update(detail)
            break

    # if abstract not found
    # if 'start' not in abstract:
    #    raise AbstractNotFoundError("Abstract header not found")

    author_paragraphs = doc.paragraphs[title_start+1: abstract['start']]

    authors = []
    for p in author_paragraphs:
        if p.text.strip():
            superscript_removed_text = ''  # remove superscript footnotes
            for r in p.runs:
                superscript_removed_text += r.text if not r.font.superscript else ''
            style_ok, detail = check_style(p, AUTHOR_DETAILS)
            author_details = {
                'text': superscript_removed_text,
                'style': p.style.name,
                'style_ok': style_ok,
            }
            author_details.update(detail)
            authors.append(author_details)
    return abstract, authors


def get_text(p):
    text = ''.join([r.text.upper() if r.style and (r.style.font.all_caps or r.font.all_caps) else r.text for r in p.runs])
    if p.style.font.all_caps or p.style.base_style and p.style.base_style.font.all_caps:
        text = text.upper()
    return text


def check_title_case(text, ratio=0.7):
    if text:
        return (sum(map(str.isupper, text)) / len(list(filter(str.isalpha, text)))) > ratio
    else:
        return False


def convert_twips_to_cm(twips):
    width = Twips(int(twips))
    return round(width.mm / 10, 2)


def check_tracking_on(doc):
    for i, p in enumerate(doc.paragraphs):
        element = p._element
        for child in element.iterchildren():
            if '}ins ' in str(child) or '}del ' in child or 'proofErr' in child:
                return True

    return False
